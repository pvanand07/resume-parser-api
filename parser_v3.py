from typing import Dict, Any, Optional
from PyPDF2 import PdfReader
import re
import tiktoken
import time
import os
import json
import requests
import hashlib
import concurrent.futures
from tqdm import tqdm
import logging
from openai import OpenAI
from os import getenv
import base64
from pathlib import Path
import fitz  # PyMuPDF
import tempfile
from geocode import geocoder  # Import the geocoder instance directly
from dotenv import load_dotenv
import docx  # Import docx for Word document processing

# Load environment variables
load_dotenv()

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Get OpenRouter API key from environment if not provided at runtime
DEFAULT_API_KEY = os.getenv("OPENROUTER_API_KEY", "")

# Initialize geocoder is now done by importing from geocode module

def calculate_file_hash(file_path):
    """Calculate MD5 hash of a file."""
    hasher = hashlib.md5()
    with open(file_path, 'rb') as f:
        for chunk in iter(lambda: f.read(4096), b''):
            hasher.update(chunk)
    return hasher.hexdigest()

def sanitize_text(text):
    """Remove sensitive information like SSNs from text, but keep phone numbers."""
    # Only remove social security numbers (US format)
    text = re.sub(r'\b\d{3}[-.]?\d{2}[-.]?\d{4}\b(?![-.]?\d)', '[SSN]', text)
    
    return text

def get_encoding():
    """Get the tokenizer encoding based on the model."""
    try:
        return tiktoken.get_encoding("cl100k_base")
    except:
        return tiktoken.get_encoding("p50k_base")

def truncate_by_tokens(text, max_tokens):
    """Truncate text to stay within token limit."""
    encoding = get_encoding()
    tokens = encoding.encode(text)
    
    if len(tokens) <= max_tokens:
        return text
    
    truncated_tokens = tokens[:max_tokens]
    return encoding.decode(truncated_tokens)

def convert_pdf_to_images(pdf_path):
    """Convert PDF pages to images using PyMuPDF."""
    try:
        # Create a local directory to store images
        pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
        output_dir = os.path.join(os.path.dirname(pdf_path), f"{pdf_name}_pages")
        os.makedirs(output_dir, exist_ok=True)
        
        # Open the PDF
        doc = fitz.Document(pdf_path)  # Using Document instead of open
        image_paths = []
        
        # Convert each page to an image with higher resolution
        zoom_x = 2.0  # horizontal zoom
        zoom_y = 2.0  # vertical zoom
        mat = fitz.Matrix(zoom_x, zoom_y)  # zoom factor 2 in each dimension
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            # Render page to an image with higher resolution
            pix = page.get_pixmap(matrix=mat)
            image_path = os.path.join(output_dir, f'page_{page_num+1}.png')
            pix.save(image_path)
            image_paths.append(image_path)
            logger.info(f"Saved page {page_num+1} to {image_path}")
        
        doc.close()
        return image_paths
    except Exception as e:
        logger.error(f"Error converting PDF to images: {str(e)}")
        return []

def encode_image_to_base64(image_path):
    """Encode image to base64 string."""
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')

def extract_text_from_image(image_path, api_key):
    """Extract text from an image using OCR."""
    if not api_key:
        logger.error("No API key provided for OCR")
        return ""
        
    try:
        # Encode image to base64
        base64_image = encode_image_to_base64(image_path)
        data_url = f"data:image/jpeg;base64,{base64_image}"

        # Initialize OpenAI client with OpenRouter
        client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key,
        )

        # Make the API call
        completion = client.chat.completions.create(
            model="google/gemini-2.5-flash-preview",
            extra_headers={
                "HTTP-Referer": "https://github.com/your-repo",  # Optional
                "X-Title": "Resume Parser",  # Optional
            },
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Extract all the text from the given image"
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": data_url
                            }
                        }
                    ]
                }
            ]
        )
        
        return completion.choices[0].message.content
    except Exception as e:
        logger.error(f"Error in OCR processing: {str(e)}")
        return ""

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file using python-docx."""
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX file: {str(e)}")
        return ""

def extract_text_from_file(file_path, api_key, max_retries=3):
    """Extract text from a file based on its extension."""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path, api_key, max_retries)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    else:
        logger.error(f"Unsupported file type: {file_extension}")
        return ""

def extract_text_from_pdf(file_path, api_key, max_retries=3):
    """Extract text from a PDF file using PyPDF2 with OCR fallback."""
    # First try PyPDF2 extraction
    for attempt in range(max_retries):
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            if text.strip():
                return text
            logger.warning(f"Attempt {attempt + 1}: No text content extracted from {file_path}")
        except Exception as e:
            logger.error(f"Attempt {attempt + 1}: Error extracting text from {file_path}: {e}")
            if attempt < max_retries - 1:
                time.sleep(2)  # Wait 2 seconds before retrying
                continue
    
    # If PyPDF2 fails, try OCR
    logger.info("PyPDF2 extraction failed, attempting OCR...")
    if not api_key:
        logger.error("Cannot perform OCR: No API key provided")
        return ""
        
    try:
        # Convert PDF to images
        image_paths = convert_pdf_to_images(file_path)
        if not image_paths:
            logger.error("Failed to convert PDF to images")
            return ""
        
        # Extract text from each image
        full_text = ""
        for image_path in image_paths:
            page_text = extract_text_from_image(image_path, api_key)
            if page_text:
                full_text += page_text + "\n"
        
        return full_text.strip()
    except Exception as e:
        logger.error(f"OCR extraction failed: {str(e)}")
        return ""

def get_all_resume_files_from_folder(folder_path):
    """Get all supported resume files (PDF, DOCX) from the specified folder."""
    resume_files = []
    for file in os.listdir(folder_path):
        file_lower = file.lower()
        if file_lower.endswith('.pdf') or file_lower.endswith('.docx'):
            resume_files.append(os.path.join(folder_path, file))
    return resume_files

def load_existing_results(output_file):
    """Load existing results from JSON file if it exists."""
    if os.path.exists(output_file):
        try:
            with open(output_file, 'r') as f:
                return json.load(f)
        except json.JSONDecodeError:
            logger.error(f"Error loading existing results from {output_file}. Starting with empty results.")
    return {}

def calculate_total_work_experience(work_experience):
    """Calculate total work experience in months from work_experience array."""
    total_months = 0
    
    for job in work_experience:
        start_date = job.get('start_date', '')
        end_date = job.get('end_date', '')
        
        if not start_date:
            continue
            
        # Parse start date
        try:
            start_year, start_month, start_day = map(int, start_date.split('-'))
        except:
            continue
            
        # Parse end date (use current date if job is current or end_date is empty)
        if job.get('is_current', False) or not end_date:
            from datetime import datetime
            current_date = datetime.now()
            end_year, end_month = current_date.year, current_date.month
        else:
            try:
                end_year, end_month, end_day = map(int, end_date.split('-'))
            except:
                continue
                
        # Calculate months
        total_months += (end_year - start_year) * 12 + (end_month - start_month)
        
    return total_months

def is_remote_location(location: str) -> bool:
    """Check if a location indicates remote/online work."""
    if not location:
        return True
    location_lower = location.lower()
    remote_indicators = [
        "remote",
        "online",
        "virtual",
        "work from home",
        "wfh",
        "anywhere",
        "global",
        "worldwide",
        "distributed",
        "telecommute"
    ]
    return any(indicator in location_lower for indicator in remote_indicators)

def add_location_coordinates(location_str, data_dict, field_name='location_coordinates'):
    """Add geocoded coordinates to data dictionary."""
    if not location_str or not location_str.strip() or is_remote_location(location_str):
        data_dict[field_name] = None
        return
        
    try:
        # Use geocoder to get coordinates
        location_result = geocoder.geocode(location_str)
        
        # Check if location_result exists and has required attributes
        if location_result and hasattr(location_result, 'latitude') and hasattr(location_result, 'longitude'):
            data_dict[field_name] = {
                'latitude': location_result.latitude,
                'longitude': location_result.longitude
            }
        else:
            data_dict[field_name] = None
    except Exception as e:
        logger.warning(f"Failed to geocode location {location_str}: {str(e)}")
        data_dict[field_name] = None

def parse_resume(file_path: str, api_key: Optional[str] = None, remove_pii: bool = False) -> Dict[str, Any]:
    """Parse a resume file and extract structured data as JSON.
    
    Args:
        file_path (str): Path to the resume file
        api_key (str, optional): API key for the resume parsing service. If None, uses env var.
        remove_pii (bool): If True, removes personal identifiable information (email, phone) from the output
    """
    # Use provided API key or fall back to the environment variable
    api_key = api_key or DEFAULT_API_KEY
    
    if not api_key:
        return {"error": "No API key provided or found in environment variables"}
    
    # Extract text from the resume file
    text = extract_text_from_file(file_path, api_key)
    if not text:
        return {"error": "Failed to extract text from the resume file"}
    
    # Apply minimal sanitization (only removes SSNs)
    text = sanitize_text(text)
    
    # Truncate if needed to stay within API limits
    text = truncate_by_tokens(text, 16000)  # Adjust token limit as needed
    
    # JSON example for the model to follow
    json_example = {
        "name": "John Doe",
        "email": "candidate@example.com",
        "phone": "+91 1234567890", 
        "github_url": "https://github.com/username",
        "linkedin_url": "https://linkedin.com/in/username",
        "location": "San Francisco, CA",
        "work_experience": [
            {
                "company": "Example Corp",
                "title": "Software Engineer",
                "start_date": "2020-01",  # YYYY-MM format
                "end_date": "",  # Empty string for current positions
                "is_current": True,  # Boolean to indicate current position
                "location": "San Francisco, CA",
                "responsibilities": [
                    "Developed frontend applications using React",
                    "Implemented RESTful APIs with Node.js"
                ],
                "skills": ["React", "Node.js", "JavaScript"]
            }
        ],
        "projects": [
            {
                "name": "Personal Website",
                "description": "A responsive portfolio website",
                "skills": ["HTML", "CSS", "JavaScript"],
                "url": "https://example.com"
            }
        ]
    }
    
    # Create the system prompt with the JSON example
    system_prompt = f"""
    Extract the candidate's profile data from the given resume and format it as JSON.
    
    Here's the expected JSON structure (note that work_experience and projects are optional):
    {json.dumps(json_example, indent=2)}
    
    Rules:
    1. If you cannot extract a specific field, use an empty string or empty array as appropriate
    2. work_experience and projects can be empty arrays if not found in the resume
    3. Extract skills corresponding to the work experience from responsibilities key 
    4. Extract skills corresponding to the projects from description key
    5. Return valid JSON only, no additional text or explanations
    6. Make sure all mandatory fields (name, email, phone, github_url, linkedin_url) are included, even if empty
    7. Use a consistent parsable date format following these rules:
       - For work experience dates, use the YYYY-MM-DD format (e.g., 2020-1-01 for January 1, 2020)
       - If month is not available, use YYYY-01-01 and if day is not available, use YYYY-MM-01
       - For current positions, use an empty string for end_date and set is_current to true
       - If the date format is unclear from the resume, make your best estimate and standardize
    8. For the name field:
       - Extract the candidate's full name from the resume
       - If multiple names are found, use the most prominent one (usually at the top of the resume)
       - If no name is found, use an empty string
    9. For location fields:
       - If the candidate location is not mentioned, use the latest location from the work experience
       - If the location indicates remote/online work (e.g., "Remote", "Online", "Work from Home"), return "Remote" as location
       - For physical locations, use the city and state/country format (e.g., "San Francisco, CA")

    Respond ONLY with the JSON data.
    """
    
    max_retries = 2  # Number of retries after initial attempt
    retry_delay = 2  # Delay in seconds between retries
    
    for attempt in range(max_retries + 1):
        try:
            # Initialize OpenAI client with OpenRouter
            client = OpenAI(
                base_url="https://openrouter.ai/api/v1",
                api_key=api_key,
            )
            
            # Make the API call
            completion = client.chat.completions.create(
                model="openai/gpt-4o-mini",
                extra_headers={
                    "HTTP-Referer": "https://github.com/your-repo",  # Optional
                    "X-Title": "Resume Parser",  # Optional
                },
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": text}
                ],
                temperature=0.1,
                response_format={"type": "json_object"}
            )
            
            # Get the content from the response
            content = completion.choices[0].message.content

            if not content:
                if attempt == max_retries:
                    return {"error": "API response contained no content"}
                logger.warning(f"Empty content on attempt {attempt + 1}, retrying...")
                time.sleep(retry_delay)
                continue
            
            # Parse the JSON response
            try:
                parsed_data = json.loads(str(content))  # Ensure content is a string
                # Calculate and add total work experience in months
                parsed_data['total_we_months'] = calculate_total_work_experience(parsed_data.get('work_experience', []))
                
                # Add location coordinates using our helper function
                if 'location' in parsed_data:
                    add_location_coordinates(parsed_data['location'], parsed_data)
                else:
                    parsed_data['location_coordinates'] = None
                
                # Add location coordinates for work experience
                for job in parsed_data.get('work_experience', []):
                    if 'location' in job:
                        add_location_coordinates(job['location'], job)
                    else:
                        job['location_coordinates'] = None
                
                # Remove PII if requested
                if remove_pii:
                    if 'name' in parsed_data:
                        name_parts = parsed_data['name'].split()
                        if len(name_parts) > 1:
                            parsed_data['name'] = name_parts[0] + ' ' + name_parts[-1][0] + '.'
                        else:
                            parsed_data['name'] = name_parts[0][0] + '.'
                    if 'email' in parsed_data and parsed_data['email'] and '@' in parsed_data['email']:
                        parsed_data['email'] = parsed_data['email'].split('@')[0][:4] + 'xxxx' + '@' + parsed_data['email'].split('@')[1]
                    if 'phone' in parsed_data and parsed_data['phone']:
                        phone = parsed_data['phone']
                        if len(phone) > 4:
                            parsed_data['phone'] = phone[:4] + ' xxxx xxxxxx'
                
                return parsed_data
            except json.JSONDecodeError:
                if attempt == max_retries:
                    return {"error": "Failed to parse JSON response from API"}
                logger.warning(f"JSON parse error on attempt {attempt + 1}, retrying...")
                time.sleep(retry_delay)
                continue
                
        except Exception as e:
            if attempt == max_retries:
                return {"error": f"API request failed after {max_retries + 1} attempts: {str(e)}"}
            logger.warning(f"API request failed on attempt {attempt + 1}, retrying... Error: {str(e)}")
            time.sleep(retry_delay)
            continue
    
    # This should never be reached, but added for type safety
    return {"error": "Unexpected error in API request handling"}

def process_resume(args):
    """Process a single resume with arguments needed for the task."""
    file_path, api_key, existing_results, remove_pii = args
    file_hash = calculate_file_hash(file_path)
    hash_prefix = file_hash[:4]
    
    # Skip if already processed
    if hash_prefix in existing_results:
        return hash_prefix, None, file_path
    
    try:
        result = parse_resume(file_path, api_key, remove_pii)
        # Add filename and full hash for reference
        result["source_file"] = os.path.basename(file_path)
        result["file_hash"] = file_hash
        return hash_prefix, result, file_path
    except Exception as e:
        logger.error(f"Error processing {file_path}: {str(e)}")
        return hash_prefix, {"error": f"Processing failed: {str(e)}"}, file_path

def process_resumes_parallel(resumes_folder, output_file, api_key=None, max_workers=10, limit=0, remove_pii=False):
    """Process multiple resumes in parallel and save results to a JSON file.
    
    Args:
        resumes_folder (str): Path to folder containing resume files (PDF, DOCX)
        output_file (str): Path to save the results JSON file
        api_key (str, optional): API key for the resume parsing service. If None, uses env var.
        max_workers (int): Maximum number of parallel workers
        limit (int): Maximum number of resumes to process. If 0, process all resumes.
        remove_pii (bool): If True, removes personal identifiable information (email, phone) from the output
    """
    # Use provided API key or fall back to the environment variable
    api_key = api_key or DEFAULT_API_KEY
    
    if not api_key:
        logger.error("No API key provided or found in environment variables")
        return {}
    
    # Ensure the folder exists
    if not os.path.exists(resumes_folder):
        logger.error(f"Resumes folder '{resumes_folder}' does not exist")
        return {}
    
    # Load existing results if any
    existing_results = load_existing_results(output_file)
    logger.info(f"Loaded {len(existing_results)} existing results from {output_file}")
    
    # Get all resume files (PDF, DOCX) from the folder
    resume_files = get_all_resume_files_from_folder(resumes_folder)
    if not resume_files:
        logger.error(f"No supported resume files found in {resumes_folder}")
        return {}
    
    # Apply limit if specified
    if limit > 0:
        resume_files = resume_files[:limit]
        logger.info(f"Limiting processing to {limit} resumes")
    
    logger.info(f"Found {len(resume_files)} resume files in {resumes_folder}")
    
    # Prepare arguments for processing
    process_args = [(file_path, api_key, existing_results, remove_pii) for file_path in resume_files]
    
    # Process resumes in parallel
    results = existing_results.copy()
    processed_count = 0
    skipped_count = 0
    
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(process_resume, args): args for args in process_args}
        
        # Create a progress bar
        with tqdm(total=len(resume_files), desc="Processing resumes") as pbar:
            for future in concurrent.futures.as_completed(futures):
                hash_prefix, result, file_path = future.result()
                
                if result is None:  # Skip if already processed
                    skipped_count += 1
                    logger.info(f"Skipped already processed file: {os.path.basename(file_path)}")
                else:
                    results[hash_prefix] = result
                    processed_count += 1
                    logger.info(f"Processed: {os.path.basename(file_path)} -> {hash_prefix}")
                    
                    # Save intermediate results after each successful processing
                    with open(output_file, 'w') as f:
                        json.dump(results, f, indent=2)
                
                pbar.update(1)
    
    # Final save of results
    with open(output_file, 'w') as f:
        json.dump(results, f, indent=2)
    
    logger.info(f"Resume processing completed. Processed: {processed_count}, Skipped: {skipped_count}, Total in output: {len(results)}")
    return results

# Example usage:
if __name__ == "__main__":
    resumes_folder = "Resumes"
    output_file = "resume_results.json"
    
    # Get API key from environment or use a default for testing
    api_key = os.getenv("OPENROUTER_API_KEY", "")
    if not api_key:
        logger.warning("API key not found in environment variables. Set OPENROUTER_API_KEY for API access.")
    
    # Process all resumes in parallel (set limit=0 to process all)
    results = process_resumes_parallel(resumes_folder, output_file, api_key, max_workers=10, limit=0, remove_pii=True)
    
    # Or process only a specific number of resumes
    # results = process_resumes_parallel(resumes_folder, output_file, api_key, max_workers=10, limit=5)